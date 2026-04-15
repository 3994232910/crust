import logging
import uuid
import json
from typing import Any
from pathlib import Path

logger = logging.getLogger(__name__)

from fastapi import APIRouter, BackgroundTasks, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlmodel import Session, select

from app.api.deps import CurrentUser, SessionDep
from app.core.db import engine
from app import crud
from app.models.forge import (
    AnnotateRequest,
    CompleteRequest,
    Forge,
    ForgeCreate,
    ForgePublic,
    ForgesPublic,
    ForgeUpdate,
    ImageTo3DRequest,
    LightAdjustRequest,
    LightAdjustWithScreenshotRequest,
    LightAutoOptimizeRequest,
    LightOptimizeRequest,
    ModelInfo,
    SummarizeRequest,
    ViewOptimizeRequest,
)
from app.models import Message
from app.core.ai_client import ChatMessage, ChatRequest, ai_client, AIFeatureDisabledError
from app.workflows.outline import outline_graph
from app.workflows.summarize import summarize_graph


def _handle_ai_error(e: Exception) -> None:
    """统一处理 AI 相关异常。"""
    if isinstance(e, AIFeatureDisabledError):
        raise HTTPException(
            status_code=503,
            detail={
                "error": "feature_disabled",
                "feature": e.feature,
                "message": e.message,
            },
        )
    raise HTTPException(status_code=502, detail=f"AI 服务错误: {str(e)}")


def _build_embed_text(forge: Forge) -> str:
    """拼接用于向量化的文本：标题 + 正文。"""
    parts = []
    if forge.title:
        parts.append(forge.title)
    if forge.content:
        parts.append(forge.content)
    return "\n".join(parts)


async def _refresh_embedding(forge: Forge, session: SessionDep) -> None:
    """生成并写入 embedding，失败时静默跳过，不影响主流程。"""
    text = _build_embed_text(forge)
    if not text.strip():
        return
    try:
        vectors = await ai_client.embed([text])
        forge.embedding = vectors[0]
        session.add(forge)
        session.commit()
    except AIFeatureDisabledError:
        pass  # 功能被禁用或未配置 API Key，静默跳过
    except Exception as e:
        logger.warning("embedding 生成失败 forge_id=%s: %s", forge.id, e)


async def _refresh_embedding_bg(forge_id: uuid.UUID) -> None:
    """后台任务版本：自行开 session，响应发出后异步执行。"""
    with Session(engine) as session:
        forge = session.get(Forge, forge_id)
        if forge is None:
            return
        await _refresh_embedding(forge, session)


def _sync_links_bg(forge_id: uuid.UUID, owner_id: uuid.UUID, content: str | None) -> None:
    """后台任务：解析 [[title]] 并同步双链表。"""
    with Session(engine) as session:
        crud.sync_forge_links(session=session, source_id=forge_id, owner_id=owner_id, content=content)


router = APIRouter(prefix="/forge", tags=["forge"])


@router.post("/image-to-3d", response_model=dict)
async def image_to_3d(
    *,
    current_user: CurrentUser,
    request: ImageTo3DRequest,
) -> Any:
    """图片转 3D 模型（Mock：直接返回假数据，不调用 Hunyuan3D API）。"""
    if not request.image_base64 and not request.image_url:
        raise HTTPException(status_code=400, detail="Either image_base64 or image_url is required")

    mock_filename = f"{uuid.uuid4()}_mock_model.glb"
    mock_url = f"/models/{current_user.id}/{mock_filename}"
    return {
        "model_url": mock_url,
        "filename": mock_filename,
        "message": "Mock: 3D model generated successfully (no real API call)",
    }


# 1. 基础路由
@router.get("/", response_model=ForgesPublic)
def read_forges(
    session: SessionDep,
    current_user: CurrentUser,
    skip: int = 0,
    limit: int = 100,
) -> Any:
    """
    Retrieve forges.
    """
    forges = crud.get_forges(session=session, owner_id=current_user.id, skip=skip, limit=limit)
    return ForgesPublic(data=forges, count=len(forges))


@router.post("/", response_model=ForgePublic)
async def create_forge(
    *,
    session: SessionDep,
    forge_in: ForgeCreate,
    current_user: CurrentUser,
    background_tasks: BackgroundTasks,
) -> Any:
    """
    Create new forge.
    """
    if forge_in.is_folder and not forge_in.title:
        forge_in.title = "nebula"
    if not forge_in.is_folder and not forge_in.title:
        forge_in.title = "nova"

    forge = crud.create_forge(session=session, forge_in=forge_in, owner_id=current_user.id)

    if not forge_in.is_folder:
        background_tasks.add_task(_refresh_embedding_bg, forge.id)
        if forge_in.content:
            background_tasks.add_task(_sync_links_bg, forge.id, current_user.id, forge_in.content)

    return forge

# 2. AI 工作流路由（静态路径，必须在 /{id} 之前）

@router.post("/summarize-stream")
async def summarize_forges_stream(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    request: SummarizeRequest,
) -> StreamingResponse:
    """流式知识梳理：先创建文件，再逐 token 流式返回 AI 内容，结束后保存并建立双向链接。"""
    from datetime import datetime

    forges = crud.get_forges_by_ids(
        session=session, forge_ids=request.forge_ids, owner_id=current_user.id
    )
    if not forges:
        raise HTTPException(status_code=404, detail="未找到指定笔记")

    titles_list = ", ".join(f.title or "无标题" for f in forges[:3])
    if len(forges) > 3:
        titles_list += f" 等{len(forges)}篇笔记"
    summary_title = f"知识梳理 - {titles_list}"
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    header = f"# {summary_title}\n\n> 自动生成于 {timestamp}\n\n"

    forge_in = ForgeCreate(title=summary_title, content=header, is_folder=False)
    new_forge = crud.create_forge(session=session, forge_in=forge_in, owner_id=current_user.id)
    new_forge_id = str(new_forge.id)
    new_forge_title = summary_title

    forge_contents = [
        {"title": f.title or "", "content": f.content or ""}
        for f in forges
    ]
    focus_hint = f"\n\n请重点关注：{request.focus}" if request.focus else ""
    nodes_text = "\n\n---\n\n".join(
        f"### {item['title']}\n{(item['content'] or '')[:1500]}"
        for item in forge_contents
    )
    prompt = (
        "请对以下多篇笔记进行知识梳理，提炼核心观点、关联关系和可行结论，"
        "输出结构化的 Markdown 总结报告（不要包含报告标题，直接从正文开始）。"
        f"{focus_hint}\n\n===== 笔记内容 =====\n\n{nodes_text}"
    )
    messages = [
        ChatMessage(
            role="system",
            content="你是专业的知识梳理助手，擅长跨文档提炼核心内容和关联关系。输出结构化 Markdown 报告，不加任何解释性前缀。",
        ),
        ChatMessage(role="user", content=prompt),
    ]

    async def event_stream():
        init_data = {
            "type": "init",
            "forge_id": new_forge_id,
            "title": new_forge_title,
            "header": header,
        }
        yield f"data: {json.dumps(init_data, ensure_ascii=False)}\n\n"

        full_body = ""
        try:
            async for token in ai_client.chat_stream(ChatRequest(messages=messages, max_tokens=3000)):
                full_body += token
                yield f"data: {json.dumps({'type': 'chunk', 'content': token}, ensure_ascii=False)}\n\n"
        except AIFeatureDisabledError as e:
            yield f"data: {json.dumps({'type': 'error', 'message': e.message}, ensure_ascii=False)}\n\n"
            return
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': 'AI 生成失败'}, ensure_ascii=False)}\n\n"
            return

        # Stream the wiki backlinks section
        for chunk in ["\n\n## 相关笔记\n\n"] + [f"- [[{f.title or '无标题'}]]\n" for f in forges]:
            full_body += chunk
            yield f"data: {json.dumps({'type': 'chunk', 'content': chunk}, ensure_ascii=False)}\n\n"

        # Save full content to DB using a new session
        full_content = header + full_body
        source_uuid = uuid.UUID(new_forge_id)
        with Session(engine) as save_session:
            db_forge = save_session.get(Forge, source_uuid)
            if db_forge:
                db_forge.content = full_content
                save_session.add(db_forge)
                save_session.commit()
                crud.sync_forge_links(
                    session=save_session,
                    source_id=source_uuid,
                    owner_id=current_user.id,
                    content=full_content,
                )

        yield f"data: {json.dumps({'type': 'done', 'forge_id': new_forge_id}, ensure_ascii=False)}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.post("/summarize", response_model=dict)
async def summarize_forges(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    request: SummarizeRequest,
) -> Any:
    """对多篇笔记进行知识梳理，返回结构化总结报告（LangGraph 工作流）。"""
    forges = crud.get_forges_by_ids(
        session=session, forge_ids=request.forge_ids, owner_id=current_user.id
    )

    if not forges:
        raise HTTPException(status_code=404, detail="未找到指定笔记")

    forge_contents = [
        {"title": f.title or "", "content": f.content or ""}
        for f in forges
    ]

    try:
        result = await summarize_graph.ainvoke({
            "forge_contents": forge_contents,
            "focus": request.focus or "",
            "summary": "",
            "new_forge_id": None,
            "session": session,
            "owner_id": current_user.id,
        })
    except AIFeatureDisabledError as e:
        _handle_ai_error(e)

    return {
        "summary": result["summary"],
        "count": len(forges),
        "new_forge_id": result.get("new_forge_id"),
    }


# 3. 文件导入路由（必须放在 /{id} 之前！）
IMPORT_ALLOWED_EXTENSIONS = {
    "pdf", "docx", "pptx", "xlsx", "xls",
    "html", "htm", "csv", "json", "xml",
    "epub", "txt", "md",
}

@router.post("/import-file", response_model=ForgePublic)
async def import_file_as_forge(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    parent_id: str | None = None,
) -> Any:
    """上传文档（PDF/DOCX/PPTX/XLSX/HTML 等），自动转为 Markdown 创建笔记。"""
    import os
    import tempfile
    from markitdown import MarkItDown

    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in IMPORT_ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型: .{ext}，支持: {', '.join(sorted(IMPORT_ALLOWED_EXTENSIONS))}",
        )

    contents = await file.read()
    if len(contents) > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件过大，最大支持 50MB")

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        tmp.write(contents)
        tmp_path = tmp.name

    try:
        md = MarkItDown()
        result = md.convert(tmp_path)
        markdown_content = result.text_content or ""
    except Exception as e:
        logger.error("markitdown 转换失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"文档转换失败: {e}") from e
    finally:
        os.unlink(tmp_path)

    title = Path(file.filename).stem
    parsed_parent_id = None
    if parent_id:
        import uuid as _uuid
        try:
            parsed_parent_id = _uuid.UUID(parent_id)
        except ValueError:
            raise HTTPException(status_code=400, detail="parent_id 格式无效")

    from app.models.forge import ForgeCreate as _ForgeCreate
    forge_in = _ForgeCreate(
        title=title,
        content=markdown_content,
        is_folder=False,
        parent_id=parsed_parent_id,
    )
    forge = crud.create_forge(session=session, forge_in=forge_in, owner_id=current_user.id)
    background_tasks.add_task(_refresh_embedding_bg, forge.id)
    return forge


# 4. 知识地图路由（必须放在 /{id} 之前！）
@router.post("/refresh-embeddings", response_model=dict)
async def refresh_all_embeddings(
    session: SessionDep,
    current_user: CurrentUser,
    background_tasks: BackgroundTasks,
) -> Any:
    """为当前用户所有尚未生成向量的笔记触发向量化（后台异步执行）。"""
    forges = crud.get_forges(session=session, owner_id=current_user.id, skip=0, limit=10000)
    pending = [f for f in forges if not f.is_folder and f.embedding is None]
    for f in pending:
        background_tasks.add_task(_refresh_embedding_bg, f.id)
    return {
        "queued": len(pending),
        "message": f"已触发 {len(pending)} 条笔记的向量化，后台处理中…",
    }


@router.get("/knowledge-map")
async def get_knowledge_map(
    session: SessionDep,
    current_user: CurrentUser,
) -> Any:
    """将用户所有已向量化的笔记降维为 3D 知识地图数据。

    流程：笔记 embedding → UMAP 2D 降维 → KDE 密度地形 → KMeans 聚类 → AI 主题标注
    返回地形网格、散点坐标和聚类标签，前端用 Three.js 渲染 3D 地图。
    """
    import asyncio

    forges = crud.get_forges(session=session, owner_id=current_user.id, skip=0, limit=10000)
    notes_only = [f for f in forges if not f.is_folder]
    forges_with_emb = [f for f in notes_only if f.embedding is not None]

    if len(forges_with_emb) < 2:
        return {
            "nodes": [],
            "terrain": {"grid_x": [], "grid_y": [], "grid_z": []},
            "clusters": [],
        }

    def _compute() -> dict:  # type: ignore[return]
        import numpy as np
        import umap as umap_lib
        from scipy.stats import gaussian_kde
        from sklearn.cluster import KMeans

        embeddings = np.array([f.embedding for f in forges_with_emb], dtype=np.float32)
        n = len(embeddings)

        # UMAP 2D 降维
        n_neighbors = min(15, n - 1)
        reducer = umap_lib.UMAP(
            n_components=2, random_state=42, n_neighbors=n_neighbors, min_dist=0.1
        )
        coords_2d = reducer.fit_transform(embeddings)
        raw_x, raw_y = coords_2d[:, 0], coords_2d[:, 1]

        # 归一化到 [-15, 15]，保持纵横比
        x_center = (raw_x.min() + raw_x.max()) / 2
        y_center = (raw_y.min() + raw_y.max()) / 2
        half_range = max(raw_x.max() - raw_x.min(), raw_y.max() - raw_y.min(), 1e-6) / 2
        scale = 15.0
        x = (raw_x - x_center) / half_range * scale
        y = (raw_y - y_center) / half_range * scale

        # KDE 密度地形（60×60 网格）
        kde = gaussian_kde(np.vstack([x, y]))
        margin = 2.0
        x_min, x_max = float(x.min()) - margin, float(x.max()) + margin
        y_min, y_max = float(y.min()) - margin, float(y.max()) + margin
        grid_res = 60
        xx, yy = np.meshgrid(
            np.linspace(x_min, x_max, grid_res),
            np.linspace(y_min, y_max, grid_res),
        )
        z_raw = kde(np.vstack([xx.ravel(), yy.ravel()])).reshape(xx.shape)
        z_global_min, z_global_max = float(z_raw.min()), float(z_raw.max())
        z_range = max(z_global_max - z_global_min, 1e-10)
        zz = (z_raw - z_global_min) / z_range * 10.0  # 归一化到 0-10

        # 散点 Z 高度
        z_scatter_raw = kde(np.vstack([x, y]))
        z_scatter = (z_scatter_raw - z_global_min) / z_range * 10.0

        # KMeans 聚类
        n_clusters = min(max(2, n // 8), 8)
        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init="auto")
        cluster_labels = kmeans.fit_predict(embeddings)

        # 构建节点
        nodes = [
            {
                "id": str(forges_with_emb[i].id),
                "title": forges_with_emb[i].title or "无标题",
                "content": (forges_with_emb[i].content or "")[:300],
                "x": float(x[i]),
                "y": float(y[i]),
                "z": float(z_scatter[i]),
                "cluster_id": int(cluster_labels[i]),
            }
            for i in range(n)
        ]

        # 各簇山峰（该簇中 z 最高的节点）
        clusters_raw = []
        for cid in range(n_clusters):
            indices = [i for i in range(n) if cluster_labels[i] == cid]
            if not indices:
                continue
            peak = max(indices, key=lambda i: z_scatter[i])
            clusters_raw.append({
                "id": cid,
                "label": f"主题{cid + 1}",
                "peak_x": float(x[peak]),
                "peak_y": float(y[peak]),
                "peak_z": float(z_scatter[peak]),
                "titles": [nodes[i]["title"] for i in indices[:5]],
            })

        return {
            "nodes": nodes,
            "terrain": {
                "grid_x": xx.tolist(),
                "grid_y": yy.tolist(),
                "grid_z": zz.tolist(),
            },
            "clusters_raw": clusters_raw,
        }

    map_data = await asyncio.to_thread(_compute)

    # AI 为每个簇生成主题标签（失败时保留默认标签）
    clusters = []
    for c in map_data["clusters_raw"]:
        label = c["label"]
        try:
            titles_str = "、".join(c["titles"])
            resp = await ai_client.chat(
                ChatRequest(
                    messages=[
                        ChatMessage(
                            role="system",
                            content="你是知识图谱专家。只输出 2-5 个字的主题词，不加任何标点或解释。",
                        ),
                        ChatMessage(
                            role="user",
                            content=f"以下笔记标题属于同一知识簇：{titles_str}\n请用 2-5 个字概括该簇的主题：",
                        ),
                    ],
                    max_tokens=20,
                )
            )
            label = resp.strip()[:12]
        except Exception:
            pass
        clusters.append(
            {
                "id": c["id"],
                "label": label,
                "peak_x": c["peak_x"],
                "peak_y": c["peak_y"],
                "peak_z": c["peak_z"],
            }
        )

    return {
        "nodes": map_data["nodes"],
        "terrain": map_data["terrain"],
        "clusters": clusters,
    }


# 5. 模型路由（必须放在 /{id} 之前！）
@router.post("/upload-model", response_model=dict)
async def upload_model(
    *,
    current_user: CurrentUser,
    files: list[UploadFile] = File(...),
) -> Any:
    """Upload 3D model files (supports folder upload with .gltf + .bin + textures)."""

    models_dir = Path("models") / str(current_user.id)
    models_dir.mkdir(parents=True, exist_ok=True)

    main_model_url = None

    for file in files:
        if not file.filename:
            continue

        filename = file.filename
        file_extension = filename.split(".")[-1].lower()

        if file_extension not in ["glb", "gltf", "bin", "png", "jpg", "jpeg", "webp", "txt"]:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {filename}"
            )

        contents = await file.read()
        if len(contents) > 100 * 1024 * 1024:
            raise HTTPException(
                status_code=400,
                detail=f"File too large: {filename}"
            )

        file_path = models_dir / filename
        file_path.parent.mkdir(parents=True, exist_ok=True)

        with open(file_path, "wb") as f:
            f.write(contents)

        if file_extension in ["glb", "gltf"] and main_model_url is None:
            main_model_url = f"/models/{current_user.id}/{filename}"

    if not main_model_url:
        raise HTTPException(status_code=400, detail="No valid model file found")

    return {
        "url": main_model_url,
        "filename": main_model_url.split("/")[-1],
        "size": len(files)
    }

@router.get("/models", response_model=list[ModelInfo])
def list_models(
    *,
    current_user: CurrentUser,
) -> Any:
    """List all uploaded 3D models."""
    models_dir = Path("models") / str(current_user.id)

    if not models_dir.exists():
        return []

    models = []
    for file_path in models_dir.glob("*"):
        if file_path.is_file():
            models.append({
                "url": f"/models/{current_user.id}/{file_path.name}",
                "filename": file_path.name,
                "size": file_path.stat().st_size,
            })

    return models

@router.delete("/models/{filename}")
def delete_model(
    *,
    current_user: CurrentUser,
    filename: str,
) -> Message:
    """Delete a 3D model file."""
    file_path = Path("models") / str(current_user.id) / filename

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Model not found")

    file_path.unlink()
    return Message(message="Model deleted successfully")

# 6. 导出路由（必须放在 /{id} 之前！）
EXPORT_FORMATS = {"md", "txt", "html", "docx", "pdf"}

@router.get("/{id}/export")
async def export_forge(
    session: SessionDep,
    id: uuid.UUID,
    current_user: CurrentUser,
    format: str = "md",
) -> StreamingResponse:
    """将笔记导出为 md / txt / html / docx / pdf。"""
    if format not in EXPORT_FORMATS:
        raise HTTPException(status_code=400, detail=f"不支持的格式: {format}，支持: {', '.join(sorted(EXPORT_FORMATS))}")

    forge = crud.get_forge(session=session, forge_id=id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    if forge.is_folder:
        raise HTTPException(status_code=400, detail="文件夹无法导出")

    title = forge.title or "untitled"
    content = forge.content or ""
    safe_title = "".join(c for c in title if c.isalnum() or c in " _-（）()[]【】").strip() or "untitled"

    if format == "md":
        data = content.encode("utf-8")
        media_type = "text/markdown; charset=utf-8"
        filename = f"{safe_title}.md"
        return StreamingResponse(
            iter([data]),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{_quote(filename)}'},
        )

    if format == "txt":
        import re
        txt = re.sub(r"#{1,6}\s*", "", content)
        txt = re.sub(r"\*\*(.+?)\*\*", r"\1", txt)
        txt = re.sub(r"\*(.+?)\*", r"\1", txt)
        txt = re.sub(r"`{1,3}[^`]*`{1,3}", "", txt)
        txt = re.sub(r"!\[.*?\]\(.*?\)", "", txt)
        txt = re.sub(r"\[(.+?)\]\(.*?\)", r"\1", txt)
        txt = re.sub(r"^\s*[-*>|]\s*", "", txt, flags=re.MULTILINE)
        data = txt.encode("utf-8")
        filename = f"{safe_title}.txt"
        return StreamingResponse(
            iter([data]),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{_quote(filename)}'},
        )

    if format == "html":
        import markdown as md_lib
        body = md_lib.markdown(content, extensions=["tables", "fenced_code", "toc"])
        html = f"""<!DOCTYPE html>
<html lang="zh"><head><meta charset="utf-8">
<title>{title}</title>
<style>body{{font-family:sans-serif;max-width:800px;margin:2em auto;line-height:1.7}}
pre{{background:#f5f5f5;padding:1em;overflow:auto}}code{{font-size:.9em}}</style>
</head><body>{body}</body></html>"""
        data = html.encode("utf-8")
        filename = f"{safe_title}.html"
        return StreamingResponse(
            iter([data]),
            media_type="text/html; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{_quote(filename)}'},
        )

    if format == "docx":
        import io
        import markdown as md_lib
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        doc.add_heading(title, level=0)
        for line in content.splitlines():
            stripped = line.lstrip()
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped == "":
                doc.add_paragraph("")
            else:
                p = doc.add_paragraph()
                # 处理粗体 **text**
                import re
                parts = re.split(r"\*\*(.+?)\*\*", stripped)
                bold = False
                for part in parts:
                    run = p.add_run(part)
                    run.bold = bold
                    bold = not bold
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"{safe_title}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{_quote(filename)}'},
        )

    if format == "pdf":
        import io
        import markdown as md_lib
        from weasyprint import HTML as WeasyprintHTML
        body = md_lib.markdown(content, extensions=["tables", "fenced_code"])
        html = f"""<!DOCTYPE html>
<html lang="zh"><head><meta charset="utf-8">
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+SC&display=swap');
body{{font-family:'Noto Sans SC',sans-serif;max-width:700px;margin:2em auto;line-height:1.8;font-size:13pt}}
h1,h2,h3{{margin-top:1.2em}}pre{{background:#f5f5f5;padding:.8em;white-space:pre-wrap}}
code{{font-size:.85em}}table{{border-collapse:collapse;width:100%}}
th,td{{border:1px solid #ccc;padding:.4em .8em}}
</style></head><body><h1>{title}</h1>{body}</body></html>"""
        buf = io.BytesIO()
        WeasyprintHTML(string=html).write_pdf(buf)
        buf.seek(0)
        filename = f"{safe_title}.pdf"
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename*=UTF-8\'\'{_quote(filename)}'},
        )

    raise HTTPException(status_code=400, detail="未知格式")  # unreachable


def _quote(s: str) -> str:
    """对文件名做 RFC 5987 百分号编码。"""
    from urllib.parse import quote
    return quote(s, safe="")


# 7. 动态 ID 路由（放在最后）
@router.get("/{id}", response_model=ForgePublic)
def read_forge(
    session: SessionDep,
    id: uuid.UUID,
    current_user: CurrentUser,
) -> Any:
    """Get forge by ID."""
    forge = crud.get_forge(session=session, forge_id=id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    return forge


@router.put("/{id}", response_model=ForgePublic)
async def update_forge(
    *,
    session: SessionDep,
    id: uuid.UUID,
    forge_in: ForgeUpdate,
    current_user: CurrentUser,
    background_tasks: BackgroundTasks,
) -> Any:
    """
    Update forge.
    """
    forge = crud.get_forge(session=session, forge_id=id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    update_data = forge_in.model_dump(exclude_unset=True)
    content_changed = "title" in update_data or "content" in update_data
    forge = crud.update_forge(session=session, db_forge=forge, forge_in=forge_in)

    if content_changed and not forge.is_folder:
        background_tasks.add_task(_refresh_embedding_bg, forge.id)
        background_tasks.add_task(_sync_links_bg, forge.id, forge.owner_id, forge.content)

    return forge


@router.get("/{id}/backlinks", response_model=ForgesPublic)
def get_forge_backlinks(
    session: SessionDep,
    id: uuid.UUID,
    current_user: CurrentUser,
) -> Any:
    """返回所有通过 [[title]] 引用了此笔记的笔记（反向链接）。"""
    forge = crud.get_forge(session=session, forge_id=id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    backlinks = crud.get_backlinks(session=session, forge_id=id, owner_id=current_user.id)
    return ForgesPublic(data=backlinks, count=len(backlinks))


@router.delete("/{id}")
def delete_forge(
    session: SessionDep,
    id: uuid.UUID,
    current_user: CurrentUser,
) -> Message:
    """
    Delete forge.
    """
    forge = crud.get_forge(session=session, forge_id=id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    crud.delete_forge(session=session, db_forge=forge)
    return Message(message="Forge deleted successfully")


@router.get("/{id}/recommend", response_model=list[ForgePublic])
async def recommend_forges(
    *,
    session: SessionDep,
    id: uuid.UUID,
    current_user: CurrentUser,
    limit: int = 5,
) -> Any:
    """基于向量相似度推荐关联笔记。"""
    forge = crud.get_forge(session=session, forge_id=id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    # 没有 embedding 则先生成
    if forge.embedding is None:
        await _refresh_embedding(forge, session)
        session.refresh(forge)

    if forge.embedding is None:
        return []

    results = session.exec(
        select(Forge)
        .where(Forge.owner_id == current_user.id)
        .where(Forge.id != id)
        .where(Forge.embedding.is_not(None))  # type: ignore[union-attr]
        .where(Forge.is_folder == False)  # noqa: E712
        .order_by(Forge.embedding.op("<=>")(forge.embedding))  # type: ignore[union-attr]
        .limit(limit)
    ).all()

    return results


@router.post("/{id}/embed", response_model=dict)
async def reembed_forge(
    *,
    session: SessionDep,
    id: uuid.UUID,
    current_user: CurrentUser,
) -> Any:
    """手动触发单个笔记的 embedding 生成（用于存量数据补全）。"""
    forge = crud.get_forge(session=session, forge_id=id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    await _refresh_embedding(forge, session)
    return {"message": "Embedding 已更新"}


@router.post("/{id}/outline", response_model=dict)
async def generate_outline(
    *,
    session: SessionDep,
    id: uuid.UUID,
    current_user: CurrentUser,
) -> Any:
    """为笔记生成结构化大纲（LangGraph 工作流）。"""
    forge = crud.get_forge(session=session, forge_id=id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")
    if forge.is_folder:
        raise HTTPException(status_code=400, detail="文件夹无法生成大纲")

    try:
        result = await outline_graph.ainvoke({
            "title": forge.title or "",
            "content": forge.content or "",
            "outline": "",
        })
    except AIFeatureDisabledError as e:
        _handle_ai_error(e)

    return {"outline": result["outline"]}


@router.post("/{id}/annotate-3d", response_model=dict)
async def annotate_3d_asset(
    *,
    session: SessionDep,
    id: uuid.UUID,
    current_user: CurrentUser,
    request: AnnotateRequest,
) -> Any:
    """对 3D 资产截图进行自动标注，返回标签和描述。"""
    forge = crud.get_forge(session=session, forge_id=id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    prompt = (
        "请分析这个 3D 模型的截图，输出 JSON 格式，包含：\n"
        '{"tags": ["标签1", "标签2", ...], "description": "简短描述", "category": "模型类别", "style": "风格特征"}\n'
        "只返回 JSON，不加其他文字。"
    )

    try:
        ai_response = await ai_client.analyze_image(request.screenshot, prompt)
    except AIFeatureDisabledError as e:
        _handle_ai_error(e)
    except HTTPException:
        raise HTTPException(status_code=502, detail="AI 标注失败")

    # 尝试解析 JSON，失败则返回原始文本
    try:
        annotation = json.loads(ai_response.strip().strip("```json").strip("```").strip())
    except json.JSONDecodeError:
        annotation = {"raw": ai_response}

    return {"annotation": annotation, "forge_id": str(id)}


@router.post("/{id}/complete")
async def complete_text(
    *,
    session: SessionDep,
    id: uuid.UUID,
    current_user: CurrentUser,
    request: CompleteRequest,
) -> StreamingResponse:
    """流式文案补全 / 代码优化，返回 SSE 事件流。"""
    forge = crud.get_forge(session=session, forge_id=id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not enough permissions")

    instruction = request.instruction or "请续写或优化以下内容，保持原有风格和语气。"
    context = f"笔记标题：{forge.title or '无标题'}\n\n{instruction}\n\n---\n\n{request.text}"

    messages = [
        ChatMessage(
            role="system",
            content="你是专业的写作助手，擅长续写、优化和扩展文本内容。直接输出结果，不加解释。",
        ),
        ChatMessage(role="user", content=context),
    ]

    async def event_stream():
        try:
            async for token in ai_client.chat_stream(ChatRequest(messages=messages)):
                yield f"data: {json.dumps({'token': token}, ensure_ascii=False)}\n\n"
        except AIFeatureDisabledError as e:
            error_data = {
                "error": "feature_disabled",
                "feature": e.feature,
                "message": e.message,
            }
            yield f"data: {json.dumps(error_data)}\n\n"
        except Exception:
            yield f"data: {json.dumps({'error': 'AI 生成失败'})}\n\n"
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@router.post("/ai-optimize-light", response_model=dict)
async def ai_optimize_light(
    *,
    current_user: CurrentUser,
    request: LightOptimizeRequest,
) -> Any:
    """AI automatically optimizes lighting for a 3D model."""

    config = {
        "ambient": 0.7,
        "hemisphere": {"skyColor": "#ffffff", "groundColor": "#555555", "intensity": 0.6},
        "directional": [
            {"position": [5, 10, 5], "intensity": 1.2, "color": "#ffffff"},
            {"position": [-5, 5, -5], "intensity": 0.7, "color": "#ffeedd"},
            {"position": [0, 8, -8], "intensity": 0.5, "color": "#ffffff"}
        ],
        "environment": "studio"
    }

    model_name = request.modelPath.lower()

    if any(keyword in model_name for keyword in ["car", "vehicle", "汽车"]):
        config["ambient"] = 0.8
        config["directional"] = [
            {"position": [8, 12, 6], "intensity": 1.5, "color": "#ffffff"},
            {"position": [-6, 6, -4], "intensity": 0.8, "color": "#e8f0ff"},
            {"position": [0, 10, -10], "intensity": 0.6, "color": "#ffffff"}
        ]
    elif any(keyword in model_name for keyword in ["character", "人物", "角色"]):
        config["ambient"] = 0.9
        config["hemisphere"]["intensity"] = 0.8
        config["directional"] = [
            {"position": [3, 8, 4], "intensity": 1.3, "color": "#fff5e6"},
            {"position": [-4, 6, 3], "intensity": 0.9, "color": "#e6f0ff"}
        ]
    elif any(keyword in model_name for keyword in ["building", "建筑", "house"]):
        config["ambient"] = 0.6
        config["directional"] = [
            {"position": [10, 15, 8], "intensity": 1.4, "color": "#fff8e7"},
            {"position": [-8, 8, -6], "intensity": 0.6, "color": "#d4e5ff"}
        ]

    return {"config": config, "message": "AI 已自动优化光照配置"}

@router.post("/ai-adjust-light", response_model=dict)
async def ai_adjust_light(
    *,
    current_user: CurrentUser,
    request: LightAdjustRequest,
) -> Any:
    """AI adjusts lighting based on user feedback."""

    current = request.currentConfig or {
        "ambient": 0.5,
        "hemisphere": {"skyColor": "#ffffff", "groundColor": "#444444", "intensity": 0.4},
        "directional": [{"position": [5, 10, 5], "intensity": 1.0, "color": "#ffffff"}],
        "environment": "studio"
    }

    feedback = request.feedback.lower()
    config = {
        "ambient": current.get("ambient", 0.5),
        "hemisphere": current.get("hemisphere", {}),
        "directional": current.get("directional", []),
        "environment": current.get("environment", "studio")
    }

    if "更亮" in feedback or "brighter" in feedback or "太暗" in feedback:
        config["ambient"] = min(current.get("ambient", 0.5) + 0.3, 1.5)
        for light in config["directional"]:
            light["intensity"] = min(light.get("intensity", 1.0) + 0.4, 2.0)

    elif "更暗" in feedback or "darker" in feedback or "太亮" in feedback:
        config["ambient"] = max(current.get("ambient", 0.5) - 0.2, 0.1)
        for light in config["directional"]:
            light["intensity"] = max(light.get("intensity", 1.0) - 0.3, 0.2)

    elif "左侧" in feedback or "left" in feedback:
        config["directional"].append({
            "position": [-8, 6, 4],
            "intensity": 1.0,
            "color": "#ffffff"
        })

    elif "右侧" in feedback or "right" in feedback:
        config["directional"].append({
            "position": [8, 6, 4],
            "intensity": 1.0,
            "color": "#ffffff"
        })

    elif "柔和" in feedback or "soft" in feedback or "生硬" in feedback:
        config["ambient"] = min(current.get("ambient", 0.5) + 0.2, 1.0)
        config["hemisphere"]["intensity"] = min(current.get("hemisphere", {}).get("intensity", 0.4) + 0.3, 1.0)
        config["environment"] = "city"

    elif "重置" in feedback or "reset" in feedback:
        config = {
            "ambient": 0.5,
            "hemisphere": {"skyColor": "#ffffff", "groundColor": "#444444", "intensity": 0.4},
            "directional": [{"position": [5, 10, 5], "intensity": 1.0, "color": "#ffffff"}],
            "environment": "studio"
        }

    return {"config": config, "message": "AI 已根据您的反馈调整光照"}

def parse_light_config(ai_response: str) -> dict:
    """解析 AI 返回的光照配置."""
    try:
        ai_response = ai_response.strip()

        # 尝试提取 JSON（兼容 ```json ... ``` 或纯 JSON）
        json_str = ai_response

        if "```" in ai_response:
            parts = ai_response.split("```")
            for part in parts:
                part = part.strip()
                if part.startswith("{") and part.endswith("}"):
                    json_str = part
                    break

        # 解析 JSON
        config = json.loads(json_str)

        # 校验 + 限制范围
        valid_config = {
            'ambient': max(0.1, min(config.get('ambient', 0.5), 2.0)),
            'hemisphere': {
                'skyColor': config.get('hemisphere', {}).get('skyColor', '#ffffff'),
                'groundColor': config.get('hemisphere', {}).get('groundColor', '#444444'),
                'intensity': max(
                    0.1,
                    min(config.get('hemisphere', {}).get('intensity', 0.4), 1.5)
                )
            },
            'directional': [],
            'environment': config.get('environment', 'studio')
        }

        for light in config.get('directional', []):
            pos = light.get('position', [0, 0, 0])
            valid_config['directional'].append({
                'position': [
                    max(-20, min(pos[0], 20)),
                    max(-20, min(pos[1], 20)),
                    max(-20, min(pos[2], 20))
                ],
                'intensity': max(0.1, min(light.get('intensity', 1.0), 3.0)),
                'color': light.get('color', '#ffffff')
            })

        return valid_config

    except (json.JSONDecodeError, KeyError, IndexError) as e:
        print(f'Failed to parse AI response: {e}')
        return {}


@router.post("/ai-auto-optimize-light", response_model=dict)
async def ai_auto_optimize_light(
    *,
    current_user: CurrentUser,
    request: LightAutoOptimizeRequest,
) -> Any:
    """AI 自主观测画面效果，迭代优化光照参数."""

    current = request.currentConfig or {
        'ambient': 0.5,
        'hemisphere': {'skyColor': '#ffffff', 'groundColor': '#444444', 'intensity': 0.4},
        'directional': [{'position': [5, 10, 5], 'intensity': 1.0, 'color': '#ffffff'}],
        'environment': 'studio'
    }

    iteration = request.iteration

    prompt = f"""你是一个专业的 3D 渲染光照优化专家。你的任务是分析 3D 模型的渲染画面，自动调整光照参数以达到最佳视觉效果。

请分析画面中的以下问题并优化：
1. 整体亮度是否合适（太暗或太亮）
2. 是否有死黑区域（完全没有光照）
3. 是否有过曝区域（亮度过高失去细节）
4. 阴影是否自然
5. 模型细节是否清晰可见
6. 材质表现是否良好

这是第 {iteration} 次光照优化迭代。

当前光照配置：
{json.dumps(current, ensure_ascii=False)}

只返回 JSON 格式的优化配置，不要其他文字：
{{
  "ambient": 环境光强度 (0.1-2.0),
  "hemisphere": {{
    "skyColor": "天空颜色 hex",
    "groundColor": "地面颜色 hex",
    "intensity": 半球光强度 (0.1-1.5)
  }},
  "directional": [
    {{
      "position": [x, y, z],
      "intensity": 强度 (0.1-3.0),
      "color": "颜色 hex"
    }}
  ],
  "environment": "环境贴图预设 (studio/city/park/dawn/dusk/night)"
}}"""

    try:
        ai_response = await ai_client.analyze_image(request.screenshot, prompt)
    except (HTTPException, AIFeatureDisabledError, Exception) as e:
        logger.error(f"AI auto optimize light failed: {e}")
        return {
            'config': current,
            'iteration': iteration,
            'shouldContinue': False,
            'message': 'AI 调用失败，保持当前配置'
        }

    new_config = parse_light_config(ai_response)

    if not new_config:
        return {
            'config': current,
            'iteration': iteration,
            'shouldContinue': False,
            'message': 'AI 解析失败，保持当前配置'
        }

    should_continue = iteration < 1

    return {
        'config': new_config,
        'iteration': iteration,
        'shouldContinue': should_continue,
        'message': f'AI 第 {iteration} 次优化完成'
    }


@router.post('/ai-adjust-light-with-screenshot', response_model=dict)
async def ai_adjust_light_with_screenshot(
    *,
    current_user: CurrentUser,
    request: LightAdjustWithScreenshotRequest,
) -> Any:
    """AI 根据用户反馈和画面效果调整光照."""

    current = request.currentConfig or {
        'ambient': 0.5,
        'hemisphere': {'skyColor': '#ffffff', 'groundColor': '#444444', 'intensity': 0.4},
        'directional': [{'position': [5, 10, 5], 'intensity': 1.0, 'color': '#ffffff'}],
        'environment': 'studio'
    }

    feedback = request.feedback

    prompt = f"""你是一个专业的 3D 渲染光照优化专家。根据用户的自然语言反馈，调整 3D 场景的光照配置。

用户反馈：{feedback}

当前配置：
{json.dumps(current, ensure_ascii=False)}

只返回 JSON 格式的光照配置，不要其他文字：
{{
  "ambient": 环境光强度 (0.1-2.0),
  "hemisphere": {{
    "skyColor": "天空颜色 hex",
    "groundColor": "地面颜色 hex",
    "intensity": 半球光强度 (0.1-1.5)
  }},
  "directional": [
    {{
      "position": [x, y, z],
      "intensity": 强度 (0.1-3.0),
      "color": "颜色 hex"
    }}
  ],
  "environment": "环境贴图预设 (studio/city/park/dawn/dusk/night)"
}}"""

    try:
        if request.screenshot:
            ai_response = await ai_client.analyze_image(request.screenshot, prompt)
        else:
            ai_response = await ai_client.chat(
                ChatRequest(messages=[ChatMessage(role="user", content=prompt)])
            )
    except (HTTPException, AIFeatureDisabledError, Exception) as e:
        logger.error(f"AI adjust light with screenshot failed: {e}")
        return {'config': current, 'message': 'AI 调用失败，保持当前配置'}

    new_config = parse_light_config(ai_response)

    if not new_config:
        return {'config': current, 'message': 'AI 解析失败，保持当前配置'}

    return {'config': new_config, 'message': 'AI 已根据反馈调整光照'}


@router.post('/ai-optimize-view', response_model=dict)
async def ai_optimize_view(
    *,
    current_user: CurrentUser,
    request: ViewOptimizeRequest,
) -> Any:
    """AI 分析截图，推荐最佳观察视角。"""

    current = request.currentCamera or {
        'position': [5, 5, 5],
        'target': [0, 0, 0]
    }

    prompt = f"""你是一个3D模型展示专家。分析当前3D模型的渲染截图，推荐能最好展现模型特征的观察视角。

要求：
1. 视角要能展现模型的主要特征和细节
2. 模型要完整可见，不被裁剪
3. 选择有立体感的斜45°左右视角，避免纯正面/侧面
4. 相机距离要合适，模型占画面60%-80%

当前相机配置：
{json.dumps(current, ensure_ascii=False)}

只返回JSON，不要其他文字：
{{
  "position": [x, y, z],
  "target": [x, y, z],
  "reason": "推荐理由（一句话）"
}}

注意：position 和 target 的每个分量范围约 -20 到 20，根据模型大小合理设置。"""

    try:
        ai_response = await ai_client.analyze_image(request.screenshot, prompt)
    except (HTTPException, AIFeatureDisabledError, Exception) as e:
        logger.error(f"AI optimize view failed: {e}")
        return {'camera': current, 'message': 'AI 调用失败，保持当前视角'}

    try:
        import re
        json_match = re.search(r'\{[\s\S]*\}', ai_response)
        if not json_match:
            raise ValueError('No JSON found')
        data = json.loads(json_match.group())

        pos = data.get('position', current['position'])
        tgt = data.get('target', current['target'])

        # 校验并夹紧数值
        def clamp_vec(v: list, lo: float = -20, hi: float = 20) -> list:
            return [max(lo, min(float(x), hi)) for x in v[:3]]

        camera = {
            'position': clamp_vec(pos),
            'target': clamp_vec(tgt),
        }
        reason = data.get('reason', '已优化视角')
        return {'camera': camera, 'message': reason}

    except Exception as e:
        logger.error(f"AI view parse failed: {e}, response: {ai_response[:200]}")
        return {'camera': current, 'message': 'AI 解析失败，保持当前视角'}


@router.post("/{forge_id}/publish-to-community", response_model=dict)
async def publish_forge_to_community(
    *,
    session: SessionDep,
    current_user: CurrentUser,
    forge_id: uuid.UUID,
    background_tasks: BackgroundTasks,
) -> Any:
    """将 Forge 笔记发布到社区"""
    forge = crud.get_forge(session=session, forge_id=forge_id)
    if not forge:
        raise HTTPException(status_code=404, detail="Forge not found")
    
    if forge.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    if not forge.title or not forge.content:
        raise HTTPException(status_code=400, detail="Forge must have title and content")
    
    from app.models.community import CommunityPostCreate
    
    post_in = CommunityPostCreate(
        title=forge.title,
        content=forge.content,
        source_forge_id=forge_id,
        is_published=True,
    )
    
    post = crud.create_community_post(
        session=session, post_in=post_in, owner_id=current_user.id
    )
    
    if post.content:
        background_tasks.add_task(_generate_post_embedding_bg, post.id, post.content[:2000])
    
    return {
        "message": "Published to community successfully",
        "post_id": str(post.id),
        "forge_id": str(forge_id),
    }


async def _generate_post_embedding_bg(post_id: uuid.UUID, content: str) -> None:
    """后台任务：生成帖子向量"""
    with Session(engine) as session:
        try:
            vectors = await ai_client.embed([content])
            post = session.get(CommunityPost, post_id)
            if post:
                post.embedding = vectors[0]
                session.add(post)
                session.commit()
        except Exception as e:
            logger.warning(f"Post embedding generation failed: {e}")
